"""P2b 端到端测试：经真实 FastAPI 路由验证「上传文件 → 问答命中」与类型守卫。"""

from __future__ import annotations

import io

import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine
from sqlalchemy.pool import StaticPool

import agents.ingestion_agent.pipeline  # noqa: F401  # 注册 ingestion 表
from agents.governance_agent.database.session import Base, get_async_session
from agents.ingestion_agent.fts import ensure_fts_table
from agents.ingestion_agent.storage import FakeStorage, get_storage
from agents.ingestion_agent.store import get_embedding_model, get_vector_store
from agents.ingestion_agent.tests.fakes import FakeEmbeddingModel, InMemoryVectorStore
from agents.router_agent.main import app


def _csv_bytes() -> bytes:
    return "客户名称,城市\n阿里巴巴,杭州\n腾讯,深圳\n".encode()


def _docx_bytes() -> bytes:
    import docx

    d = docx.Document()
    d.add_paragraph("阿里巴巴 总部位于杭州。")
    d.add_paragraph("腾讯科技 总部位于深圳。")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@pytest_asyncio.fixture
async def client():
    engine = create_async_engine(
        "sqlite+aiosqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    await ensure_fts_table(engine)
    factory = async_sessionmaker(engine, expire_on_commit=False)

    fake_vs = InMemoryVectorStore()
    fake_em = FakeEmbeddingModel()

    async def override_get_session():
        async with factory() as s:
            yield s

    app.dependency_overrides[get_async_session] = override_get_session
    app.dependency_overrides[get_vector_store] = lambda: fake_vs
    app.dependency_overrides[get_embedding_model] = lambda: fake_em
    app.dependency_overrides[get_storage] = lambda: FakeStorage()

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c

    app.dependency_overrides.clear()
    await engine.dispose()


@pytest.mark.asyncio
async def test_upload_csv_and_ask_hits(client: AsyncClient) -> None:
    resp = await client.post(
        "/ingest/upload",
        files={"file": ("customers.csv", _csv_bytes(), "text/csv")},
        data={"doc_type": "sales"},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["canonical"] == 2
    assert body["chunks"] == 2

    ask = await client.post("/api/data/ask", json={"query": "杭州的客户", "top_k": 3})
    assert ask.status_code == 200, ask.text
    ans = ask.json()
    assert ans["count"] >= 1
    assert "杭州" in ans["answer"]


@pytest.mark.asyncio
async def test_upload_docx_and_ask_hits(client: AsyncClient) -> None:
    resp = await client.post(
        "/ingest/upload",
        files={
            "file": (
                "contract.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"doc_type": "doc"},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["filename"] == "contract.docx"
    assert body["canonical"] >= 1

    ask = await client.post("/api/data/ask", json={"query": "杭州", "top_k": 3})
    assert ask.status_code == 200, ask.text
    ans = ask.json()
    assert ans["count"] >= 1
    assert "杭州" in ans["answer"]


@pytest.mark.asyncio
async def test_upload_rejects_unsupported_ext(client: AsyncClient) -> None:
    resp = await client.post(
        "/ingest/upload",
        files={"file": ("note.txt", b"hello", "text/plain")},
        data={"doc_type": "x"},
    )
    assert resp.status_code == 400


@pytest.mark.asyncio
async def test_upload_rejects_oversized(client: AsyncClient) -> None:
    big = b"x" * (21 * 1024 * 1024)  # > 20MB
    resp = await client.post(
        "/ingest/upload",
        files={"file": ("big.pdf", big, "application/pdf")},
        data={"doc_type": "x"},
    )
    assert resp.status_code == 413
