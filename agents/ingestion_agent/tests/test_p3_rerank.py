"""P3 端到端回归：QueryService.ask 集成查询改写 + 重排，仍能正确命中。

验证在引入改写/重排后，原有检索问答行为不被破坏，且关键词查询更稳。
"""

from __future__ import annotations

import io

import docx
import pytest

from agents.ingestion_agent.pipeline import IngestionPipeline
from agents.ingestion_agent.query import QueryService
from agents.ingestion_agent.tests.fakes import FakeEmbeddingModel, InMemoryVectorStore


@pytest.mark.asyncio
async def test_ask_with_rewrite_and_rerank_hits_target(
    session, fake_vs: InMemoryVectorStore, fake_em: FakeEmbeddingModel
) -> None:
    d = docx.Document()
    d.add_paragraph("阿里巴巴 总部 位于 杭州 市 余杭区。")  # target
    d.add_paragraph("杭州 今日 天气 多云 气温 适中。")  # distracter（含 杭州 但无关总部）
    buf = io.BytesIO()
    d.save(buf)
    data = buf.getvalue()

    await IngestionPipeline.ingest_file(
        "corp.docx", data, doc_type="corp", session=session,
        vector_store=fake_vs, embedding_model=fake_em,
    )

    ans = await QueryService.ask(
        "阿里巴巴 总部 在哪", top_k=2, doc_type="corp",
        vector_store=fake_vs, embedding_model=fake_em,
    )
    assert ans["count"] >= 1
    assert "阿里巴巴" in ans["answer"]
    assert "杭州" in ans["answer"]
    # 命中来源应包含「总部」相关文本
    assert "总部" in ans["sources"][0].get("text", "") or "总部" in ans["sources"][0].get("parent_text", "")


@pytest.mark.asyncio
async def test_ask_rewrite_removes_stopwords_but_keeps_intent(
    session, fake_vs: InMemoryVectorStore, fake_em: FakeEmbeddingModel
) -> None:
    """含停用词的口语化查询（「请问 杭州 的 客户 是 哪些」）仍应命中杭州客户文档。"""

    rows = "客户名称,城市\n阿里巴巴,杭州\n腾讯,深圳\n"
    await IngestionPipeline.ingest_file(
        "cust.csv", rows.encode("utf-8"), doc_type="cust", session=session,
        vector_store=fake_vs, embedding_model=fake_em,
    )
    ans = await QueryService.ask(
        "请问 杭州 的 客户 是 哪些", top_k=3, doc_type="cust",
        vector_store=fake_vs, embedding_model=fake_em,
    )
    assert "杭州" in ans["answer"]
    assert "阿里巴巴" in ans["answer"]
