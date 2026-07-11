"""P2b 文件入库流水线测试：ingest_file（csv）→ DB + Qdrant + 检索命中。"""

from __future__ import annotations

import pytest
from sqlalchemy import select

from agents.ingestion_agent.database.models import (
    CanonicalDocument as CanonicalDocumentORM,
)
from agents.ingestion_agent.database.models import DocumentChunk
from agents.ingestion_agent.pipeline import IngestionPipeline
from agents.ingestion_agent.query import QueryService
from agents.ingestion_agent.tests.fakes import FakeEmbeddingModel, InMemoryVectorStore


def _csv_bytes() -> bytes:
    return "客户名称,Order-No,城市\n阿里巴巴,SO-1,杭州\n腾讯,SO-2,深圳\n".encode()


@pytest.mark.asyncio
async def test_ingest_file_csv_stores_canonical_and_chunks(
    session, fake_vs: InMemoryVectorStore, fake_em: FakeEmbeddingModel
) -> None:
    result = await IngestionPipeline.ingest_file(
        "customers.csv",
        _csv_bytes(),
        doc_type="sales",
        session=session,
        vector_store=fake_vs,
        embedding_model=fake_em,
    )
    assert result["canonical"] == 2
    assert result["chunks"] == 2
    assert result["indexed_vectors"] == 2

    canon = (await session.execute(select(CanonicalDocumentORM))).scalars().all()
    assert len(canon) == 2
    chunks = (await session.execute(select(DocumentChunk))).scalars().all()
    assert len(chunks) == 2
    assert fake_vs.points  # 向量已写入内存库


@pytest.mark.asyncio
async def test_ingest_file_then_ask_hits(
    session, fake_vs: InMemoryVectorStore, fake_em: FakeEmbeddingModel
) -> None:
    await IngestionPipeline.ingest_file(
        "customers.csv",
        _csv_bytes(),
        doc_type="sales",
        session=session,
        vector_store=fake_vs,
        embedding_model=fake_em,
    )
    ans = await QueryService.ask(
        "杭州的客户",
        top_k=3,
        doc_type="sales",
        vector_store=fake_vs,
        embedding_model=fake_em,
    )
    assert ans["count"] >= 1
    assert "杭州" in ans["answer"]
    assert "阿里巴巴" in ans["answer"]
    # 父子 chunk：命中应回带父块（整表）上下文
    assert "腾讯" in ans["sources"][0]["parent_text"]


@pytest.mark.asyncio
async def test_ingest_file_docx_then_ask_hits(
    session, fake_vs: InMemoryVectorStore, fake_em: FakeEmbeddingModel
) -> None:
    import docx

    d = docx.Document()
    d.add_heading("合同清单", level=1)
    d.add_paragraph("阿里巴巴 总部位于杭州。")
    d.add_paragraph("腾讯科技 总部位于深圳。")
    import io

    buf = io.BytesIO()
    d.save(buf)
    data = buf.getvalue()

    result = await IngestionPipeline.ingest_file(
        "contract.docx",
        data,
        doc_type="doc",
        session=session,
        vector_store=fake_vs,
        embedding_model=fake_em,
    )
    assert result["canonical"] >= 1
    assert result["chunks"] >= 1

    ans = await QueryService.ask(
        "杭州",
        top_k=3,
        doc_type="doc",
        vector_store=fake_vs,
        embedding_model=fake_em,
    )
    assert ans["count"] >= 1
    assert "杭州" in ans["answer"]
